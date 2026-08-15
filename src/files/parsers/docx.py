"""Bounded parser for Microsoft Word (.docx)."""

from src.core.config import settings


class DocxParser:
    @staticmethod
    def parse_docx(file_path: str) -> tuple[str | None, bool]:
        try:
            import docx

            doc = docx.Document(file_path)
            content_parts: list[str] = []
            total_chars = 0
            limit = settings.max_document_extract_chars

            def add(text: str) -> bool:
                nonlocal total_chars
                text = text.strip()
                if not text:
                    return total_chars >= limit
                remaining = max(0, limit - total_chars)
                content_parts.append(text[:remaining])
                total_chars += min(len(text), remaining)
                return total_chars >= limit

            for paragraph in doc.paragraphs:
                if add(paragraph.text):
                    break
            if total_chars < limit:
                for table in doc.tables:
                    stop = False
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        if add(row_text):
                            stop = True
                            break
                    if stop:
                        break

            full_text = "\n".join(content_parts).strip()
            return full_text if full_text else None, True
        except Exception as exc:
            return f"Error parsing DOCX: {exc}", False


docx_parser = DocxParser()
